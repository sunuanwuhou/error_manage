#!/usr/bin/env python3
"""
docx_to_md.py
将行测题 docx 转换为 Markdown：
  - 数量关系图片（第34、35、41题）→ OCR 识别文字
  - 图形推理 / 科学推理 / 资料分析图片 → 直接嵌入 <img>
"""

import os
import re
import io
import sys
import json
import argparse
import urllib.request
import urllib.error
import numpy as np
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from PIL import Image, ImageEnhance

# ── 配置 ──────────────────────────────────────────
_DEFAULT_DOCX = "/Users/10030299/Documents/个人/2025年广东省公务员录用考试《行测》题（网友回忆版）.docx"
OUT_DIR   = Path(__file__).parent / "output"
IMG_DIR   = OUT_DIR / "images"
OCR_URL   = "http://localhost:8000/ocr/recognize"

# 章节关键词 → 内部标识
SECTION_KEYWORDS = {
    "政治理论": "politics",
    "常识判断": "common",
    "言语理解": "verbal",   # 国考/联考有此章节
    "数量关系": "math",
    "判断推理": "reasoning",
    "科学推理": "science",
    "资料分析": "data",
}

# 需要 OCR 识别的章节（数量关系）
OCR_SECTIONS = {"math"}

# ── OCR 调用 ─────────────────────────────────────

def ocr_image_bytes(img_bytes: bytes, filename: str = "img.png", captcha_mode: bool = False) -> str:
    """调用本地 OCR 服务，返回识别文本（失败时返回空串）"""
    boundary = "----DocxMdBoundary"
    body = (
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="file"; filename="{filename}"\r\n'
        f"Content-Type: image/png\r\n\r\n"
    ).encode() + img_bytes + (
        f"\r\n--{boundary}\r\n"
        f'Content-Disposition: form-data; name="captcha_mode"\r\n\r\n{"true" if captcha_mode else "false"}'
        f"\r\n--{boundary}\r\n"
        f'Content-Disposition: form-data; name="return_confidence"\r\n\r\nfalse'
        f"\r\n--{boundary}--\r\n"
    ).encode()

    req = urllib.request.Request(
        OCR_URL, data=body,
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            result = json.loads(resp.read())
            return result.get("data", {}).get("full_text", "")
    except Exception as e:
        print(f"  [OCR 失败] {e}", file=sys.stderr)
        return ""


# ── OCR 结果后处理 ────────────────────────────────

# 全角数字 → 半角
_FULLWIDTH_TABLE = str.maketrans('０１２３４５６７８９', '0123456789')
# 数学场景下常见误识别字符（·=1 出现在分子位置）
_OCR_FIXES = str.maketrans({'了': '7', 'Ｏ': '0', 'Ｚ': '2', 'Ｓ': '5', 'Ｉ': '1', '·': '1'})


# 竖排分数：分割线被识别成 /l/ 或 /I/，分母里 ↓ 是 4 的误读
_VERT_FRAC = re.compile(r'^(\d+)/[lI]/(.+)$')
_VERT_FRAC_DENOM_FIXES = str.maketrans({'↓': '4', '↑': '1'})


def normalize_math_ocr(text: str) -> str:
    """全角数字转半角 + 数字常见误识别修正（仅用于数量关系图片）"""
    text = text.translate(_FULLWIDTH_TABLE)
    text = text.translate(_OCR_FIXES)
    # 竖排分数修复：\d+/I/<denom> → \d+/<denom>（若分母修正后全为数字）
    m = _VERT_FRAC.match(text)
    if m:
        denom = m.group(2).translate(_VERT_FRAC_DENOM_FIXES)
        if re.match(r'^\d+$', denom):
            text = f"{m.group(1)}/{denom}"
    return text


# OCR 结果中出现以下字符，说明识别已乱码，直接回退到嵌图
_OCR_GARBAGE_CHARS = set(r'\}$↓')


def ocr_result_ok(text: str) -> bool:
    """粗略判断 OCR 结果是否可信（True=可用文字，False=回退到图片）"""
    if not text:
        return False
    if any(c in text for c in _OCR_GARBAGE_CHARS):
        return False
    return True


# ── OCR 前预处理 ──────────────────────────────────

def preprocess_for_ocr(img_bytes: bytes) -> bytes:
    """OCR 前预处理：小图等比放大（最短边 < 200px）+ 增强对比度"""
    try:
        img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
        w, h = img.size
        min_side = min(w, h)
        if min_side < 200:
            scale = 200 / min_side
            img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
        img = ImageEnhance.Contrast(img).enhance(2.0)
        buf = io.BytesIO()
        img.save(buf, "PNG")
        return buf.getvalue()
    except Exception:
        return img_bytes


# ── 图片提取 ─────────────────────────────────────

def get_para_images(para, doc_part) -> list[tuple[str, bytes]]:
    """提取段落内所有图片，返回 [(rId, blob), ...]"""
    result = []
    seen = set()
    for blip in para._p.findall('.//' + qn('a:blip')):
        rId = blip.get(qn('r:embed'))
        if rId and rId not in seen and rId in doc_part.rels:
            rel = doc_part.rels[rId]
            if 'image' in rel.reltype:
                result.append((rId, rel.target_part.blob))
                seen.add(rId)
    return result



# ── 段落文字清洗 ─────────────────────────────────

def clean_text(text: str) -> str:
    # 去掉行末多余空格，保留内容
    return text.rstrip()


def is_separator(text: str) -> bool:
    return bool(re.fullmatch(r'[-─=]{5,}', text.strip()))


def is_answer_line(text: str) -> bool:
    return text.strip().startswith("正确答案:")


# ── 主转换逻辑 ────────────────────────────────────

def convert(docx_path: Path, out_dir: Path, img_dir: Path, md_path: Path):
    out_dir.mkdir(parents=True, exist_ok=True)
    doc = Document(docx_path)
    doc_part = doc.part

    lines_md: list[str] = []
    current_section = "politics"
    img_counter = 0
    # 记录已处理的 rId，避免同一图片重复保存
    saved_rids: dict[str, str] = {}

    total = len(doc.paragraphs)
    print(f"共 {total} 段落，开始处理...")

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # ── 章节切换检测 ──
        for kw, sec in SECTION_KEYWORDS.items():
            if text.startswith(kw) and len(text) < 120:
                current_section = sec
                print(f"  [段落{idx}] 进入章节: {kw} → {sec}")
                break

        # ── 分隔线：转为 md 水平线 ──
        if is_separator(text):
            lines_md.append("\n---\n")
            continue

        # ── 答案行：加引用块标记 ──
        if is_answer_line(text):
            lines_md.append(f"\n> {text}\n")
            continue

        # ── 提取图片 ──
        images = get_para_images(para, doc_part)
        img_tags: list[str] = []

        for rId, blob in images:
            # 已保存过的同一张图直接复用文件名
            if rId in saved_rids:
                img_filename = saved_rids[rId]
            else:
                img_counter += 1
                img_filename = f"img_{img_counter:03d}.png"
                # 统一转为 PNG（RGBA/透明背景 → 白底合成，避免透明变黑）
                try:
                    pil = Image.open(io.BytesIO(blob))
                    if pil.mode in ('RGBA', 'LA', 'P'):
                        pil = pil.convert("RGBA")
                        bg = Image.new("RGB", pil.size, (255, 255, 255))
                        bg.paste(pil, mask=pil.split()[3])
                        pil = bg
                    else:
                        pil = pil.convert("RGB")
                    buf = io.BytesIO()
                    pil.save(buf, "PNG")
                    img_dir.mkdir(parents=True, exist_ok=True)
                    (img_dir / img_filename).write_bytes(buf.getvalue())
                    saved_rids[rId] = img_filename
                    print(f"  [段落{idx}] 保存图片: {img_filename}  章节={current_section}")
                except Exception as e:
                    print(f"  [段落{idx}] 图片保存失败: {e}", file=sys.stderr)
                    continue

            img_path_rel = f"images/{img_filename}"

            if current_section in OCR_SECTIONS:
                orig_img = Image.open(img_dir / img_filename)
                # 宽度 < 200px 的小图（行内分数/公式）才 OCR，大图（数阵/表格）直接嵌入
                if orig_img.width < 200:
                    blob_png = (img_dir / img_filename).read_bytes()
                    is_small = min(orig_img.width, orig_img.height) < 100
                    blob_for_ocr = preprocess_for_ocr(blob_png)
                    ocr_text = normalize_math_ocr(
                        ocr_image_bytes(blob_for_ocr, img_filename, captcha_mode=is_small)
                    )
                    print(f"  [OCR] {img_filename} → {ocr_text!r}")
                    if ocr_result_ok(ocr_text):
                        img_tags.append(f"`{ocr_text}`")
                    else:
                        img_tags.append(f"![]({img_path_rel})")
                else:
                    # 大图直接嵌入
                    print(f"  [大图嵌入] {img_filename} ({orig_img.width}x{orig_img.height})")
                    img_tags.append(f"![]({img_path_rel})")
            else:
                # 图推 / 科学 / 资料 → 直接嵌入
                img_tags.append(f"![]({img_path_rel})")

        # ── 拼合文字 + 图片标签 ──
        parts = []
        if text:
            parts.append(clean_text(text))
        parts.extend(img_tags)

        if parts:
            line = "  ".join(parts)

            # 题目行加粗题号
            m = re.match(r'^(\d+)\.\s', line)
            if m:
                num = m.group(1)
                line = f"**{num}.** " + line[len(m.group(0)):]

            # 章节标题加 ##
            for kw in SECTION_KEYWORDS:
                if text.startswith(kw) and len(text) < 120:
                    line = f"## {line}"
                    break

            lines_md.append(line)
        elif images:
            # 纯图片段落（无文字）
            lines_md.extend(img_tags)

    # ── 写入 MD ──
    md_content = "\n".join(lines_md)
    md_path.write_text(md_content, encoding="utf-8")
    print(f"\n✅ 完成！MD 文件: {md_path}")
    print(f"   图片目录: {img_dir}  ({img_counter} 张)")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="将行测题 docx 转换为 Markdown")
    parser.add_argument("docx", nargs="?", default=_DEFAULT_DOCX, help="docx 文件路径")
    parser.add_argument("-o", "--out", help="输出目录（默认: ./output/<docx文件名>/）")
    args = parser.parse_args()

    docx_path = Path(args.docx)
    if not docx_path.exists():
        print(f"❌ 文件不存在: {docx_path}", file=sys.stderr)
        sys.exit(1)

    # 按 docx 文件名区分输出目录
    if args.out:
        OUT_DIR = Path(args.out)
    else:
        OUT_DIR = Path(__file__).parent / "output" / docx_path.stem
    IMG_DIR = OUT_DIR / "images"
    MD_PATH = OUT_DIR / (docx_path.stem + ".md")

    convert(docx_path, OUT_DIR, IMG_DIR, MD_PATH)
